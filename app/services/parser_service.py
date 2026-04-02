"""
Parser Service — handles extraction of text from PDF, DOCX, TXT, CSV files.
"""

import csv
import io
import logging
from pathlib import Path
from typing import Optional

import PyPDF2
import docx
import pandas as pd

from app.schemas.documents import ParsedDocument, ParsedSection, FileType

logger = logging.getLogger(__name__)


class ParserService:
    """Extracts raw text and structure from supported file types."""

    # ── PDF ──────────────────────────────────────────────
    def parse_pdf(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        text_parts: list[str] = []
        sections: list[ParsedSection] = []

        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
                    sections.append(ParsedSection(
                        heading=f"Page {i + 1}",
                        content=page_text.strip(),
                        section_index=i,
                    ))
        except Exception as e:
            logger.error(f"PDF parse error ({filename}): {e}")
            text_parts = [f"[Error parsing PDF: {e}]"]

        return ParsedDocument(
            document_id=doc_id,
            filename=filename,
            file_type=FileType.PDF,
            raw_text="\n\n".join(text_parts),
            sections=sections,
        )

    # ── DOCX ─────────────────────────────────────────────
    def parse_docx(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        sections: list[ParsedSection] = []
        text_parts: list[str] = []

        try:
            doc = docx.Document(file_path)
            current_heading: Optional[str] = None
            current_content: list[str] = []
            idx = 0

            for para in doc.paragraphs:
                if para.style and para.style.name.startswith("Heading"):
                    # Flush previous section
                    if current_content:
                        sections.append(ParsedSection(
                            heading=current_heading,
                            content="\n".join(current_content),
                            section_index=idx,
                        ))
                        idx += 1
                        current_content = []
                    current_heading = para.text.strip()
                else:
                    if para.text.strip():
                        current_content.append(para.text.strip())

                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Flush last section
            if current_content:
                sections.append(ParsedSection(
                    heading=current_heading,
                    content="\n".join(current_content),
                    section_index=idx,
                ))
        except Exception as e:
            logger.error(f"DOCX parse error ({filename}): {e}")
            text_parts = [f"[Error parsing DOCX: {e}]"]

        return ParsedDocument(
            document_id=doc_id,
            filename=filename,
            file_type=FileType.DOCX,
            raw_text="\n\n".join(text_parts),
            sections=sections,
        )

    # ── TXT ──────────────────────────────────────────────
    def parse_txt(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        try:
            text = Path(file_path).read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            logger.error(f"TXT parse error ({filename}): {e}")
            text = f"[Error reading TXT: {e}]"

        return ParsedDocument(
            document_id=doc_id,
            filename=filename,
            file_type=FileType.TXT,
            raw_text=text,
            sections=[ParsedSection(content=text, section_index=0)],
        )

    # ── CSV ──────────────────────────────────────────────
    def parse_csv(self, file_path: str, doc_id: str, filename: str) -> ParsedDocument:
        try:
            df = pd.read_csv(file_path)
            text = df.to_string(index=False)
            columns = list(df.columns)
            row_count = len(df)

            # Also create a markdown-like summary
            summary = f"Columns: {', '.join(columns)}\nRows: {row_count}\n\n"
            summary += df.head(50).to_csv(index=False)
        except Exception as e:
            logger.error(f"CSV parse error ({filename}): {e}")
            text = f"[Error parsing CSV: {e}]"
            columns = []
            row_count = 0
            summary = text

        return ParsedDocument(
            document_id=doc_id,
            filename=filename,
            file_type=FileType.CSV,
            raw_text=summary,
            sections=[ParsedSection(content=summary, section_index=0)],
            row_count=row_count,
            column_names=columns,
        )

    # ── Router ───────────────────────────────────────────
    def parse(self, file_path: str, doc_id: str, filename: str, file_type: FileType) -> ParsedDocument:
        """Route to the correct parser based on file type."""
        parsers = {
            FileType.PDF: self.parse_pdf,
            FileType.DOCX: self.parse_docx,
            FileType.TXT: self.parse_txt,
            FileType.CSV: self.parse_csv,
        }
        parser_fn = parsers.get(file_type)
        if not parser_fn:
            raise ValueError(f"Unsupported file type: {file_type}")
        return parser_fn(file_path, doc_id, filename)


# Singleton
parser_service = ParserService()
